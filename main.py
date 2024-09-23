import os
from PyQt6.QtWidgets import QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, QListWidget, QFileDialog, QMessageBox, QInputDialog
from docx import Document
from openpyxl import load_workbook

def find_text_in_docx(file_path, search_text):
    doc = Document(file_path)
    search_result = []
    for para in doc.paragraphs:
        if search_text.lower() in para.text.lower():
            search_result.append((file_path, para.text))
    for table in doc.tables:
        for i in table.rows:
            for j in i.cells:
                if search_text.lower() in j.text.lower():
                    search_result.append((file_path, j.text))
    return search_result

def find_text_in_xlsx(file_path, search_text):
    wb = load_workbook(file_path, data_only=True)
    search_result = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if isinstance(cell, str) and search_text.lower() in cell.lower():
                    search_result.append((file_path, cell))
    return search_result

def edit_text_in_docx(file_path, old_text, new_text):
    doc = Document(file_path)
    for para in doc.paragraphs:
        if old_text in para.text:
            para.text = para.text.replace(old_text, new_text)
    for table in doc.tables:
        for i in table.rows:
            for j in i.cells:
                if old_text in j.text:
                    j.text = j.text.replace(old_text, new_text)
    doc.save(file_path)

def edit_text_in_xlsx(file_path, old_text, new_text):
    wb = load_workbook(file_path)
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and old_text in cell.value:
                    cell.value = cell.value.replace(old_text, new_text)
    wb.save(file_path)

def search_files(dir, search_text):
    search_result = []
    for i in os.listdir(dir):
        if i.endswith(".docx"):
            file_path = os.path.join(dir, i)
            search_result.extend(find_text_in_docx(file_path, search_text))
        elif i.endswith(".xlsx"):
            file_path = os.path.join(dir, i)
            search_result.extend(find_text_in_xlsx(file_path, search_text))
    return search_result

class Documents(QWidget):
    def __init__(self):
        super().__init__()
        self.gUI()

    def gUI(self):
        self.setWindowTitle('Редактор документов')
        self.setGeometry(100, 100, 800, 600)

        layout = QVBoxLayout()

        dir_layout = QHBoxLayout()
        dir_label = QLabel('Выберите директорию:')
        self.dir_input = QLineEdit()
        dir_button = QPushButton('Обзор')
        dir_button.clicked.connect(self.select_dir)
        dir_layout.addWidget(dir_label)
        dir_layout.addWidget(self.dir_input)
        dir_layout.addWidget(dir_button)

        search_layout = QHBoxLayout()
        search_label = QLabel('Текст для поиска:')
        self.search_input = QLineEdit()
        search_button = QPushButton('Поиск')
        search_button.clicked.connect(self.search_text)
        search_layout.addWidget(search_label)
        search_layout.addWidget(self.search_input)
        search_layout.addWidget(search_button)

        search_result_layout = QVBoxLayout()
        search_result_label = QLabel('Найденные строки:')
        self.search_result_list = QListWidget()
        edit_button = QPushButton('Редактировать')
        edit_button.clicked.connect(self.edit_selected_text)
        search_result_layout.addWidget(search_result_label)
        search_result_layout.addWidget(self.search_result_list)
        search_result_layout.addWidget(edit_button)

        layout.addLayout(dir_layout)
        layout.addLayout(search_layout)
        layout.addLayout(search_result_layout)

        self.setLayout(layout)

    def select_dir(self):
        dir = QFileDialog.getExistingDirectory(self, 'Выберите директорию')
        if dir:
            self.dir_input.setText(dir)

    def search_text(self):
        dir = self.dir_input.text()
        search_text = self.search_input.text()
        if not dir or not search_text:
            QMessageBox.critical(self, 'Ошибка', 'Пожалуйста, заполните все поля!')
        else:
            search_result = search_files(dir, search_text)
            self.search_result_list.clear()
            for i in search_result:
                self.search_result_list.addItem(f"{i[0]}: {i[1]}")

    def edit_selected_text(self):
        selected_item = self.search_result_list.currentItem()
        if selected_item:
            selected_match = selected_item.text()
            file_path, old_text = selected_match.split(': ', 1)
            new_text, ok = QInputDialog.getText(self, 'Редактирование', f'Введите новый текст для:\n{old_text}', text=old_text)
            if ok and new_text:
                if file_path.endswith(".docx"):
                    edit_text_in_docx(file_path, old_text, new_text)
                elif file_path.endswith(".xlsx"):
                    edit_text_in_xlsx(file_path, old_text, new_text)
                QMessageBox.information(self, 'Успех', 'Текст успешно изменен!')
                self.search_text()
        else:
            QMessageBox.critical(self, 'Ошибка', 'Пожалуйста, выберите строку для редактирования.')

if __name__ == '__main__':
    app = QApplication([])
    ex = Documents()
    ex.show()
    app.exec()
